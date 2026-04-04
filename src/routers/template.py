import ast
import html
import json
import threading
from pathlib import Path
from typing import List, Optional, Union
from uuid import uuid4

from fastapi import APIRouter, HTTPException, Depends, BackgroundTasks
from pydantic import BaseModel, Field
from prompt import (
    TEMPLATE_ANALYSE_PROMPT,
    TEMPLATE_BUILD_USER_PROMPT_TEMPLATE,
    TEMPLATE_DESCRIPTION_USER_PROMPT_TEMPLATE,
    TEMPLATE_RAG_SUMMARY_SYSTEM_PROMPT,
    TEMPLATE_RAG_SUMMARY_USER_PROMPT_TEMPLATE,
)
from sqlalchemy.orm import Session
from sqlalchemy import or_
from database import get_db, SessionLocal
from models import Template, User, Log
from graphrag.graphrag_service import get_graphrag_service, GRAPHRAG_IMPORT_ERROR
from routers.user import get_current_user
from utils.model import ask_messages, LLMError

router = APIRouter(tags=["Template"])

SERVER_ROOT_DIR = Path(__file__).resolve().parents[2]
SUMMARY_EXPORT_DIR = Path(__file__).resolve().parents[2] / "analysis_results" / "summary_exports"
SUMMARY_JOB_STORE: dict[str, dict] = {}
SUMMARY_JOB_LOCK = threading.Lock()

SUMMARY_PROGRESS = {
    "queued": 5,
    "validating": 15,
    "searching": 35,
    "building_prompt": 55,
    "generating": 78,
    "exporting": 92,
    "completed": 100,
    "failed": 100,
}

SUMMARY_STATUS_MESSAGE = {
    "queued": "任务已创建，等待执行",
    "validating": "校验模板和参数中",
    "searching": "正在检索相关文档片段",
    "building_prompt": "正在构建摘要上下文",
    "generating": "正在调用大模型生成摘要",
    "exporting": "正在导出结果文件",
    "completed": "摘要生成完成",
    "failed": "摘要生成失败",
}

class TemplateResponse(BaseModel):
    content: dict


class TemplateRequest(BaseModel):
    """保存模板请求模型"""
    name: str
    prompt: str
    category: int = 0
    description: str = None
    example: str = None
    icon_path: str = None
    labels: Optional[Union[List[str], str]] = None


class TemplateUpdateRequest(BaseModel):
    """更新模板请求模型"""
    name: Optional[str] = None
    prompt: Optional[str] = None
    category: Optional[int] = None
    description: Optional[str] = None
    example: Optional[str] = None
    icon_path: Optional[str] = None
    labels: Optional[Union[List[str], str]] = None


class TemplateDuplicateRequest(BaseModel):
    """复制模板请求模型"""
    name: Optional[str] = None


class TemplateSummaryRequest(BaseModel):
    """基于模板 + GraphRAG 生成摘要请求模型。"""
    query_text: str = Field(..., description="用户关注主题/问题")
    paper_ids: Optional[List[int]] = Field(default=None, description="可选：限定论文 ID 范围")
    top_k: int = Field(default=8, ge=1, le=30, description="向量召回数量")
    focus_direction: str = Field(default="行业发展趋势与技术演进路径", description="摘要聚焦方向")
    style: str = Field(default="客观严谨的学术/商业报告风格", description="文风要求")
    word_limit: str = Field(default="500-800字", description="篇幅要求")
    graph_top_entities: int = Field(default=8, ge=1, le=20, description="每篇论文抽取的核心实体数")
    snippets_per_entity: int = Field(default=2, ge=1, le=5, description="每个实体保留的上下文片段数")
    neighbor_limit: int = Field(default=4, ge=0, le=20, description="图谱实体邻居数量")
    max_graph_papers: int = Field(default=3, ge=1, le=10, description="最多补充图谱上下文的论文数")
    use_concept_expansion: bool = Field(default=True, description="是否启用 Query+CSO 概念扩展摘要链路")
    concept_max_hops: int = Field(default=2, ge=1, le=2, description="Concept 拓展最大跳数，建议 1-2")


def _update_summary_job(job_id: str, **fields) -> None:
    with SUMMARY_JOB_LOCK:
        job = SUMMARY_JOB_STORE.get(job_id)
        if not job:
            return
        job.update(fields)


def _update_summary_job_stage(job_id: str, stage: str) -> None:
    _update_summary_job(
        job_id,
        status="running" if stage not in {"completed", "failed"} else stage,
        stage=stage,
        progress=SUMMARY_PROGRESS.get(stage, 0),
        message=SUMMARY_STATUS_MESSAGE.get(stage, ""),
    )


def _execute_summary_generation(
    template_id: int,
    request: TemplateSummaryRequest,
    current_user: User,
    db: Session,
    job_id: Optional[str] = None,
) -> dict:
    if job_id:
        _update_summary_job_stage(job_id, "validating")

    template = (
        db.query(Template)
        .filter(
            Template.id == template_id,
            or_(Template.user_id == current_user.id, Template.user_id == 0),
        )
        .first()
    )
    if template is None:
        raise HTTPException(status_code=404, detail="模板不存在或无权限访问")

    if GRAPHRAG_IMPORT_ERROR is not None:
        raise HTTPException(
            status_code=500,
            detail=f"neo4j-graphrag 导入失败: {GRAPHRAG_IMPORT_ERROR}",
        )

    if job_id:
        _update_summary_job_stage(job_id, "searching")

    graphrag_service = get_graphrag_service()
    search_result = graphrag_service.similarity_search(
        request.query_text,
        top_k=request.top_k,
        paper_ids=request.paper_ids,
        strict_paper_filter=True,
    )
    chunks = search_result.get("results", [])
    if not chunks:
        payload = {
            "code": 200,
            "message": "未检索到相关片段，无法生成摘要",
            "data": {
                "template_id": template.id,
                "template_name": template.name,
                "query_text": request.query_text,
                "summary": "未检索到相关片段，请调整检索问题或放宽 paper_ids 范围。",
                "retrieved_chunks": 0,
                "graph_context_blocks": 0,
            },
        }
        if job_id:
            _update_summary_job_stage(job_id, "completed")
            _update_summary_job(job_id, result=payload)
        return payload

    if job_id:
        _update_summary_job_stage(job_id, "building_prompt")

    citations, paper_citation_ids = _build_citation_mapping(chunks)
    document_chunks = _format_document_chunks(chunks, paper_citation_ids)
    graph_relations, concept_context_meta = _format_graph_context(graphrag_service, request, chunks)
    if not graph_relations.strip():
        graph_relations = "暂无可用图谱上下文。"

    prompt_text = TEMPLATE_RAG_SUMMARY_USER_PROMPT_TEMPLATE.format(
        template_prompt=template.prompt,
        query_text=request.query_text,
        focus_direction=request.focus_direction,
        style=request.style,
        word_limit=request.word_limit,
        document_chunks=document_chunks,
        graph_relations=graph_relations,
    )
    prompt_text += (
        "\n\n输出格式要求：\n"
        "1) 必须输出标准 Markdown。\n"
        "2) 结构需包含标题、分节小标题和要点列表。\n"
        "3) 引用请保留 [文献引用标记：x]，不要输出 JSON。"
    )

    if job_id:
        _update_summary_job_stage(job_id, "generating")

    llm_result = ask_messages(
        messages=[
            {"role": "system", "content": TEMPLATE_RAG_SUMMARY_SYSTEM_PROMPT},
            {"role": "user", "content": prompt_text},
        ],
        max_tokens=2200,
        temperature=0.2,
        top_p=0.8,
        extra_payload={
            "skip_special_tokens": False,
            "spaces_between_special_tokens": False,
            "chat_template_kwargs": {"enable_thinking": False},
        },
    )

    summary_markdown = (llm_result.content or "").strip()
    if not summary_markdown:
        raise HTTPException(status_code=502, detail="大模型返回为空，未生成有效 Markdown")

    if job_id:
        _update_summary_job_stage(job_id, "exporting")

    result_id = str(uuid4())
    export_info = _export_summary_bundle(summary_markdown, result_id)
    export_info["provenance_json_path"] = _export_provenance_debug_json(
        result_id=result_id,
        query_text=request.query_text,
        chunks=chunks,
        paper_citation_ids=paper_citation_ids,
    )

    graph_block_count = graph_relations.count("【实体】")
    knowledge_ids = _build_knowledge_ids_for_log(search_result, chunks)
    log_entry = Log(
        user_id=current_user.id,
        template_id=template.id,
        knowledge_ids=knowledge_ids,
        result_path=export_info["markdown_path"],
    )
    db.add(log_entry)
    db.commit()
    db.refresh(log_entry)

    payload = {
        "code": 200,
        "message": "摘要生成成功",
        "data": {
            "template_id": template.id,
            "template_name": template.name,
            "query_text": request.query_text,
            "result_id": result_id,
            "summary": summary_markdown,
            "summary_markdown": summary_markdown,
            "retrieved_chunks": len(chunks),
            "graph_context_blocks": graph_block_count,
            "concept_expanded_topics": concept_context_meta.get("expanded_topics", []),
            "concept_anchor_entities": concept_context_meta.get("anchor_entities", []),
            "paper_ids": search_result.get("paper_ids"),
            "citations": citations,
            "log_id": log_entry.id,
            "files": export_info,
        },
    }
    if job_id:
        _update_summary_job_stage(job_id, "completed")
        _update_summary_job(job_id, result=payload)
    return payload


def _run_summary_job(job_id: str, template_id: int, request_payload: dict, user_id: int) -> None:
    db = SessionLocal()
    try:
        user = db.query(User).filter(User.id == user_id).first()
        if user is None:
            raise HTTPException(status_code=404, detail="用户不存在，无法执行摘要任务")

        request = TemplateSummaryRequest(**request_payload)
        _execute_summary_generation(
            template_id=template_id,
            request=request,
            current_user=user,
            db=db,
            job_id=job_id,
        )
    except HTTPException as exc:
        _update_summary_job_stage(job_id, "failed")
        _update_summary_job(job_id, error=str(exc.detail))
    except LLMError as exc:
        _update_summary_job_stage(job_id, "failed")
        _update_summary_job(job_id, error=f"无法连接到大模型服务: {exc}")
    except Exception as exc:
        _update_summary_job_stage(job_id, "failed")
        _update_summary_job(job_id, error=f"基于模板生成摘要失败: {exc}")
    finally:
        db.close()

def get_template_prompt(description: str) -> str:
    return TEMPLATE_DESCRIPTION_USER_PROMPT_TEMPLATE.format(description=description)


def extract_first_brace_block(text: str) -> str:
    """Return substring from first '{' to last '}' if both exist."""
    start = text.find("{")
    end = text.rfind("}")
    if start == -1 or end == -1 or start >= end:
        return text
    return text[start:end + 1]


def _normalize_labels(labels: Optional[Union[List[str], str]]) -> Optional[str]:
    """将标签输入统一转换为逗号分隔字符串，便于存储到 templates.labels。"""
    if labels is None:
        return None

    if isinstance(labels, list):
        cleaned = [str(item).strip() for item in labels if str(item).strip()]
        return ",".join(cleaned) if cleaned else None

    cleaned = [item.strip() for item in str(labels).split(",") if item.strip()]
    return ",".join(cleaned) if cleaned else None


def _split_labels(labels: Optional[str]) -> List[str]:
    """将数据库中的逗号分隔 labels 转回前端使用的 tags 数组。"""
    if not labels:
        return []
    return [item.strip() for item in labels.split(",") if item.strip()]


def _format_document_chunks(chunks: List[dict], paper_citation_ids: dict) -> str:
    """将召回片段格式化为可读上下文，并按文献分配唯一引用编号。"""
    lines: List[str] = []
    for idx, chunk in enumerate(chunks, start=1):
        score = chunk.get("score")
        score_text = f"{float(score):.4f}" if score is not None else "N/A"
        paper_id = chunk.get("paper_id")
        paper_key = f"paper:{paper_id}" if paper_id is not None else f"title:{(chunk.get('title') or '').strip()}"
        citation_id = paper_citation_ids.get(paper_key)
        lines.append(
            (
                f"[Chunk {idx}] [文献引用标记：{citation_id}]"
                # f"年份={chunk.get('year')}，chunk_index={chunk.get('chunk_index')}，相似度={score_text}\n"
                f"{chunk.get('text') or ''}"
            )
        )
    print("\n".join(lines))
    return "\n\n".join(lines)


def _build_citation_mapping(chunks: List[dict]) -> tuple[dict, dict]:
    """按文献维度构建引用映射，保证一篇文献仅对应一个引用标记。"""
    citations: dict = {}
    paper_citation_ids: dict = {}

    for chunk in chunks:
        paper_id = chunk.get("paper_id")
        title = (chunk.get("title") or "未知").strip() or "未知"
        paper_key = f"paper:{paper_id}" if paper_id is not None else f"title:{title}"

        if paper_key not in paper_citation_ids:
            citation_id = len(paper_citation_ids) + 1
            paper_citation_ids[paper_key] = citation_id
            citations[citation_id] = {
                "id": citation_id,
                "paper_id": paper_id,
                "title": title,
                "year": chunk.get("year"),
                "chunk_ids": [],
                "scores": [],
                "text_excerpts": [],
            }

        citation_id = paper_citation_ids[paper_key]
        citation_item = citations[citation_id]
        chunk_id = chunk.get("chunk_id")
        if chunk_id and chunk_id not in citation_item["chunk_ids"]:
            citation_item["chunk_ids"].append(chunk_id)

        score = chunk.get("score")
        if score is not None:
            citation_item["scores"].append(float(score))

        excerpt = (chunk.get("text") or "")[:200]
        if excerpt and excerpt not in citation_item["text_excerpts"] and len(citation_item["text_excerpts"]) < 3:
            citation_item["text_excerpts"].append(excerpt)

    for citation in citations.values():
        scores = citation.pop("scores", [])
        citation["max_score"] = max(scores) if scores else None
        citation["avg_score"] = (sum(scores) / len(scores)) if scores else None

    return citations, paper_citation_ids


def _format_graph_context(service, request: TemplateSummaryRequest, chunks: List[dict]) -> tuple[str, dict]:
    """基于召回论文补充图谱上下文；优先接入 query+concept 扩展链路。"""
    paper_ids: List[int] = []
    for chunk in chunks:
        paper_id = chunk.get("paper_id")
        if isinstance(paper_id, int) and paper_id not in paper_ids:
            paper_ids.append(paper_id)

    target_paper_ids = request.paper_ids or paper_ids[: request.max_graph_papers]
    concept_context_meta: dict = {
        "expanded_topics": [],
        "anchor_entities": [],
    }

    graph_context_parts: List[str] = []

    if request.use_concept_expansion and request.query_text.strip() and target_paper_ids:
        try:
            concept_result = service.generate_query_concept_summary(
                query_text=request.query_text,
                paper_ids=target_paper_ids,
                anchor_entity_limit=request.graph_top_entities,
                concept_max_hops=request.concept_max_hops,
                direct_snippets_per_paper=request.snippets_per_entity * 3,
                expanded_snippets_per_paper=request.snippets_per_entity * 3,
            )
            concept_context_meta["expanded_topics"] = concept_result.get("expanded_topics", [])
            concept_context_meta["anchor_entities"] = [
                {
                    "name": item.get("name"),
                    "type": item.get("type"),
                    "score": item.get("score"),
                }
                for item in concept_result.get("anchor_entities", [])
            ]

            expanded_topics = concept_result.get("expanded_topics", [])
            if expanded_topics:
                graph_context_parts.append("[Concept 扩展主题]\n" + "、".join(expanded_topics))

            papers = concept_result.get("papers", [])
            for paper in papers:
                direct_chunks = paper.get("direct_chunks", [])
                expanded_chunks = paper.get("expanded_chunks", [])
                direct_preview = "\n".join([f"- {c[:220]}" for c in direct_chunks[:4]]) or "- 无"
                expanded_preview = "\n".join([f"- {c[:220]}" for c in expanded_chunks[:4]]) or "- 无"
                graph_context_parts.append(
                    "\n".join(
                        [
                            f"[Paper {paper.get('paper_id')}] {paper.get('title') or '未知标题'}",
                            f"直接主题: {', '.join(paper.get('direct_topics') or []) or '无'}",
                            f"扩展主题: {', '.join(paper.get('expanded_topics') or []) or '无'}",
                            "直接证据:",
                            direct_preview,
                            "关联延伸证据:",
                            expanded_preview,
                        ]
                    )
                )

            concept_summary = (concept_result.get("summary") or "").strip()
            if concept_summary:
                graph_context_parts.append("[Concept 融合综述草案]\n" + concept_summary)
        except Exception:
            # 概念链路异常时，自动回退到原图谱上下文逻辑。
            pass

    if not graph_context_parts:
        for paper_id in paper_ids[: request.max_graph_papers]:
            try:
                refined = service.get_graph_refined_context(
                    paper_id=paper_id,
                    top_entities=request.graph_top_entities,
                    snippets_per_entity=request.snippets_per_entity,
                    neighbor_limit=request.neighbor_limit,
                )
            except Exception:
                continue

            blocks = refined.get("context_blocks", [])
            if blocks:
                graph_context_parts.append(f"[Paper {paper_id}]\n" + "\n\n".join(blocks))

    return "\n\n".join(graph_context_parts), concept_context_meta


def _build_provenance_debug_records(chunks: List[dict], paper_citation_ids: dict) -> List[dict]:
    """构建溯源调试记录，保留片段内容和文献名。"""
    records: List[dict] = []
    for idx, chunk in enumerate(chunks, start=1):
        paper_id = chunk.get("paper_id")
        title = (chunk.get("title") or "未知").strip() or "未知"
        paper_key = f"paper:{paper_id}" if paper_id is not None else f"title:{title}"
        score = chunk.get("score")

        records.append(
            {
                "rank": idx,
                "citation_id": paper_citation_ids.get(paper_key),
                "paper_id": paper_id,
                "paper_title": title,
                "chunk_id": chunk.get("chunk_id"),
                "chunk_index": chunk.get("chunk_index"),
                "score": float(score) if score is not None else None,
                "snippet_text": chunk.get("text") or "",
            }
        )

    return records


def _build_knowledge_ids_for_log(search_result: dict, chunks: List[dict]) -> str:
    """为日志构建 knowledge_ids 字段，优先使用检索返回的 paper_ids。"""
    raw_ids = search_result.get("paper_ids") or []
    unique_ids: List[str] = []

    for item in raw_ids:
        if isinstance(item, int):
            item_str = str(item)
            if item_str not in unique_ids:
                unique_ids.append(item_str)

    if not unique_ids:
        for chunk in chunks:
            paper_id = chunk.get("paper_id")
            if isinstance(paper_id, int):
                paper_id_str = str(paper_id)
                if paper_id_str not in unique_ids:
                    unique_ids.append(paper_id_str)

    # logs.knowledge_ids 为 VARCHAR(255)，超长时进行截断避免写库失败。
    return ",".join(unique_ids)[:255]


def _ensure_summary_export_dir() -> Path:
    SUMMARY_EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    return SUMMARY_EXPORT_DIR


def _export_markdown_file(markdown_text: str, file_path: Path) -> None:
    file_path.write_text(markdown_text, encoding="utf-8")


def _export_docx_file(markdown_text: str, file_path: Path) -> None:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("缺少 python-docx 依赖，无法导出 Word") from exc

    document = Document()
    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip()
        if not line:
            document.add_paragraph("")
            continue

        if line.startswith("### "):
            document.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            document.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            document.add_heading(line[2:].strip(), level=1)
        elif line.startswith("- "):
            document.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            document.add_paragraph(line)

    document.save(file_path)


def _export_pdf_file(markdown_text: str, file_path: Path) -> None:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
    except ImportError as exc:
        raise RuntimeError("缺少 reportlab 依赖，无法导出 PDF") from exc

    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))

    doc = SimpleDocTemplate(
        str(file_path),
        pagesize=A4,
        leftMargin=48,
        rightMargin=48,
        topMargin=54,
        bottomMargin=54,
    )

    styles = getSampleStyleSheet()
    body_style = ParagraphStyle(
        "BodyCN",
        parent=styles["Normal"],
        fontName="STSong-Light",
        fontSize=11,
        leading=18,
    )
    h1_style = ParagraphStyle(
        "H1CN",
        parent=styles["Heading1"],
        fontName="STSong-Light",
        fontSize=18,
        leading=24,
    )
    h2_style = ParagraphStyle(
        "H2CN",
        parent=styles["Heading2"],
        fontName="STSong-Light",
        fontSize=15,
        leading=21,
    )
    h3_style = ParagraphStyle(
        "H3CN",
        parent=styles["Heading3"],
        fontName="STSong-Light",
        fontSize=13,
        leading=18,
    )

    story = []
    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip()
        if not line:
            story.append(Spacer(1, 8))
            continue

        escaped_line = html.escape(line)
        if line.startswith("### "):
            story.append(Paragraph(html.escape(line[4:].strip()), h3_style))
        elif line.startswith("## "):
            story.append(Paragraph(html.escape(line[3:].strip()), h2_style))
        elif line.startswith("# "):
            story.append(Paragraph(html.escape(line[2:].strip()), h1_style))
        elif line.startswith("- "):
            bullet_text = html.escape(line[2:].strip())
            story.append(Paragraph(f"• {bullet_text}", body_style))
        else:
            story.append(Paragraph(escaped_line, body_style))

    if not story:
        story.append(Paragraph("(Empty Markdown)", body_style))

    doc.build(story)


def _export_summary_bundle(markdown_text: str, result_id: str) -> dict:
    export_dir = _ensure_summary_export_dir()
    md_path = export_dir / f"{result_id}.md"
    docx_path = export_dir / f"{result_id}.docx"
    pdf_path = export_dir / f"{result_id}.pdf"

    _export_markdown_file(markdown_text, md_path)
    _export_docx_file(markdown_text, docx_path)
    _export_pdf_file(markdown_text, pdf_path)

    def _to_server_relative(path: Path) -> str:
        try:
            return path.relative_to(SERVER_ROOT_DIR).as_posix()
        except ValueError:
            # 异常场景回退，避免接口直接返回 Windows 绝对盘符。
            return path.name

    return {
        "result_id": result_id,
        "output_dir": _to_server_relative(export_dir),
        "markdown_path": _to_server_relative(md_path),
        "word_path": _to_server_relative(docx_path),
        "pdf_path": _to_server_relative(pdf_path),
    }


def _export_provenance_debug_json(
    result_id: str,
    query_text: str,
    chunks: List[dict],
    paper_citation_ids: dict,
) -> str:
    """导出溯源调试 JSON 文件。"""
    export_dir = _ensure_summary_export_dir()
    json_path = export_dir / f"{result_id}_provenance.json"

    payload = {
        "result_id": result_id,
        "query_text": query_text,
        "records": _build_provenance_debug_records(chunks, paper_citation_ids),
    }
    json_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")

    try:
        return json_path.relative_to(SERVER_ROOT_DIR).as_posix()
    except ValueError:
        return json_path.name

@router.post("/template/build", response_model=TemplateResponse)
def build_template(request: str):
    """
    接收用户描述，调用大模型生成模板
    """
    # 构造 Prompt，根据你的模型特性进行微调
    messages = [
        {"role": "system", "content": TEMPLATE_ANALYSE_PROMPT},
        {
            "role": "user",
            "content": TEMPLATE_BUILD_USER_PROMPT_TEMPLATE.format(text=request),
        }
    ]

    try:
        llm_result = ask_messages(
            messages=messages,
            max_tokens=4096,
            temperature=0.95,
            top_p=0.6,
            extra_payload={
                "skip_special_tokens": False,
                "spaces_between_special_tokens": False,
                "chat_template_kwargs": {"enable_thinking": False},
            },
        )
        content = llm_result.content or ""

        content = extract_first_brace_block(content)
        try:
            content_dict = ast.literal_eval(content)
        except Exception:
            content_dict = {}
        
        return TemplateResponse(content=content_dict)
    except LLMError as exc:
        raise HTTPException(
            status_code=503,
            detail=f"无法连接到大模型服务: {exc}"
        )
    except Exception as exc:
        raise HTTPException(
            status_code=500,
            detail=f"模板生成过程中发生错误: {exc}"
        )


@router.put("/template/add", status_code=200)
def add_template(information: TemplateRequest, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    """
    将生成的模板保存到数据库
    """
    try:
        # 创建新的 Template 对象，user_id 从 token 中提取
        new_template = Template(
            user_id=current_user.id,
            name=information.name,
            prompt=information.prompt,
            category=information.category,
            description=information.description,
            example=information.example,
            icon_path=information.icon_path,
            labels=_normalize_labels(information.labels),
        )
        
        db.add(new_template)
        db.commit()
        db.refresh(new_template)
        
        return {
            "code": 200,
            "message": "模板保存成功",
            "data": {
                "template_id": new_template.id,
                "name": new_template.name,
                "created_at": new_template.created_at.isoformat() if new_template.created_at else None
            }
        }
    
    except Exception as e:
        db.rollback()
        raise HTTPException(
            status_code=400,
            detail=f"保存模板失败: {str(e)}"
        )


@router.put("/template/{template_id}", status_code=200)
def update_template(
    template_id: int,
    information: TemplateUpdateRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """更新当前用户自己的模板"""
    template = (
        db.query(Template)
        .filter(Template.id == template_id, Template.user_id == current_user.id)
        .first()
    )
    if template is None:
        raise HTTPException(status_code=404, detail="模板不存在或无权限修改")

    try:
        if information.name is not None:
            template.name = information.name
        if information.prompt is not None:
            template.prompt = information.prompt
        if information.category is not None:
            template.category = information.category
        if information.description is not None:
            template.description = information.description
        if information.example is not None:
            template.example = information.example
        if information.icon_path is not None:
            template.icon_path = information.icon_path
        if information.labels is not None:
            template.labels = _normalize_labels(information.labels)

        db.commit()
        db.refresh(template)

        return {
            "code": 200,
            "message": "模板更新成功",
            "data": {
                "id": template.id,
                "user_id": template.user_id,
                "name": template.name,
                "prompt": template.prompt,
                "category": template.category,
                "description": template.description,
                "example": template.example,
                "icon_path": template.icon_path,
                "labels": template.labels,
                "tags": _split_labels(template.labels),
                "created_at": template.created_at.isoformat() if template.created_at else None,
                "updated_at": template.updated_at.isoformat() if template.updated_at else None,
            },
        }
    except Exception as exc:
        db.rollback()
        raise HTTPException(status_code=400, detail=f"更新模板失败: {str(exc)}")


@router.delete("/template/{template_id}", status_code=200)
def delete_template(
    template_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """删除当前用户自己的模板"""
    template = (
        db.query(Template)
        .filter(Template.id == template_id, Template.user_id == current_user.id)
        .first()
    )
    if template is None:
        raise HTTPException(status_code=404, detail="模板不存在或无权限删除")

    try:
        db.delete(template)
        db.commit()
        return {
            "code": 200,
            "message": "模板删除成功",
            "data": {"id": template_id},
        }
    except Exception as exc:
        db.rollback()
        raise HTTPException(status_code=400, detail=f"删除模板失败: {str(exc)}")


@router.post("/template/{template_id}/duplicate", status_code=200)
def duplicate_template(
    template_id: int,
    payload: TemplateDuplicateRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """复制用户可见模板（本人模板或公共模板）到当前用户名下"""
    source_template = (
        db.query(Template)
        .filter(
            Template.id == template_id,
            or_(Template.user_id == current_user.id, Template.user_id == 0),
        )
        .first()
    )
    if source_template is None:
        raise HTTPException(status_code=404, detail="模板不存在或无权限复制")

    try:
        duplicate_name = (payload.name or "").strip()
        if not duplicate_name:
            duplicate_name = f"{source_template.name} 副本"

        new_template = Template(
            user_id=current_user.id,
            name=duplicate_name,
            prompt=source_template.prompt,
            category=source_template.category,
            description=source_template.description,
            example=source_template.example,
            icon_path=source_template.icon_path,
            labels=source_template.labels,
        )
        db.add(new_template)
        db.commit()
        db.refresh(new_template)

        return {
            "code": 200,
            "message": "模板复制成功",
            "data": {
                "id": new_template.id,
                "user_id": new_template.user_id,
                "name": new_template.name,
                "prompt": new_template.prompt,
                "category": new_template.category,
                "description": new_template.description,
                "example": new_template.example,
                "icon_path": new_template.icon_path,
                "labels": new_template.labels,
                "tags": _split_labels(new_template.labels),
                "created_at": new_template.created_at.isoformat() if new_template.created_at else None,
                "updated_at": new_template.updated_at.isoformat() if new_template.updated_at else None,
            },
        }
    except Exception as exc:
        db.rollback()
        raise HTTPException(status_code=400, detail=f"复制模板失败: {str(exc)}")


@router.post("/template/{template_id}/summary", status_code=200)
def generate_summary_by_template(
    template_id: int,
    request: TemplateSummaryRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """根据模板 ID + GraphRAG 检索结果生成结构化全局摘要。"""
    try:
        return _execute_summary_generation(
            template_id=template_id,
            request=request,
            current_user=current_user,
            db=db,
        )
    except LLMError as exc:
        raise HTTPException(status_code=503, detail=f"无法连接到大模型服务: {exc}")
    except Exception as exc:
        import traceback; traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"基于模板生成摘要失败: {exc}")


@router.post("/template/{template_id}/summary-jobs", status_code=200)
def create_summary_job(
    template_id: int,
    request: TemplateSummaryRequest,
    background_tasks: BackgroundTasks,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """创建摘要生成任务，前端可通过 job_id 轮询进度。"""
    template = (
        db.query(Template)
        .filter(
            Template.id == template_id,
            or_(Template.user_id == current_user.id, Template.user_id == 0),
        )
        .first()
    )
    if template is None:
        raise HTTPException(status_code=404, detail="模板不存在或无权限访问")

    job_id = str(uuid4())
    with SUMMARY_JOB_LOCK:
        SUMMARY_JOB_STORE[job_id] = {
            "job_id": job_id,
            "user_id": current_user.id,
            "template_id": template_id,
            "status": "queued",
            "stage": "queued",
            "progress": SUMMARY_PROGRESS["queued"],
            "message": SUMMARY_STATUS_MESSAGE["queued"],
            "result": None,
            "error": None,
        }

    background_tasks.add_task(
        _run_summary_job,
        job_id,
        template_id,
        request.dict(),
        current_user.id,
    )

    return {
        "code": 200,
        "message": "摘要任务创建成功",
        "data": {
            "job_id": job_id,
            "status": "queued",
            "stage": "queued",
            "progress": SUMMARY_PROGRESS["queued"],
            "message": SUMMARY_STATUS_MESSAGE["queued"],
        },
    }


@router.get("/template/summary-jobs/{job_id}", status_code=200)
def get_summary_job_status(
    job_id: str,
    current_user: User = Depends(get_current_user),
):
    """查询摘要任务状态。"""
    with SUMMARY_JOB_LOCK:
        job = SUMMARY_JOB_STORE.get(job_id)

    if job is None or job.get("user_id") != current_user.id:
        raise HTTPException(status_code=404, detail="摘要任务不存在")

    return {
        "code": 200,
        "message": "查询成功",
        "data": {
            "job_id": job["job_id"],
            "status": job["status"],
            "stage": job["stage"],
            "progress": job["progress"],
            "message": job.get("message"),
            "result": job.get("result"),
            "error": job.get("error"),
        },
    }


@router.get("/template/my", status_code=200)
def get_my_templates(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    获取当前用户可见模板：
    1. 用户自己的模板（user_id = 当前用户id）
    2. 公共模板（user_id = 0）
    """
    try:
        templates = (
            db.query(Template)
            .filter(or_(Template.user_id == current_user.id, Template.user_id == 0))
            .order_by(Template.created_at.desc())
            .all()
        )

        return {
            "code": 200,
            "message": "获取模板成功",
            "data": [
                {
                    "id": template.id,
                    "user_id": template.user_id,
                    "name": template.name,
                    "prompt": template.prompt,
                    "category": template.category,
                    "description": template.description,
                    "example": template.example,
                    "icon_path": template.icon_path,
                    "labels": template.labels,
                    "tags": _split_labels(template.labels),
                    "created_at": template.created_at.isoformat() if template.created_at else None,
                    "updated_at": template.updated_at.isoformat() if template.updated_at else None,
                }
                for template in templates
            ],
        }
    except Exception as exc:
        raise HTTPException(
            status_code=500,
            detail=f"获取模板失败: {str(exc)}",
        )